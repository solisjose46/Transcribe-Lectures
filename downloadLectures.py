import docx
import extract_subtitles

# For extracting text from lecture videos.
# Can extract from Youtube and Wistia videos if they have captions.
# Can extract from multiple videos and save them to a summary document.
# Place video urls in a file called ./lectures.txt and run the script.
# Video urls should be in the following format:

# "TITLE"
# "URL"

# To find video url for youtube videos:
# Right click on video and click "Copy video URL"

# To find video url for wistia videos:
# Go to Network tab and filter with "m38u" and search for .m38u file, check response and copy url on #EXT-X-MEDIA:TYPE=SUBTITLES line
# should have value like: https://fast.wistia.net/embed/captions/n8em8nhchj.m3u8?language=eng

# or get caption url from developer console:
# Go to developer console and run the following code: document.querySelectorAll("track")
# and copy the src attribute for each track element
# Should have value like: https://fast.wistia.net/embed/captions/3ijia54m8c.vtt?language=eng

# TODO: Add command line arguments to specify input and output files names and paths

if __name__ == "__main__":
    summary_doc = docx.Document()

    # Create a file called collectedLectures.txt and write the title and video id of each lecture to the file
    with open('lectures.txt', 'r') as lectures:
        for line in lectures:
            title = line
            # remove whitespaces from title and format file name
            title = title.replace(" ", "_")
            file_name = "./lectures/{title}.docx".format(title=title)

            # move to next line
            video_url = next(lectures)
            # remove whitespaces from url
            video_url = video_url.strip()

            # Get subtitles from url
            paragraph = ""
            try:
                paragraph = extract_subtitles.get_subtitles_from_url(video_url=video_url)
            except Exception as e:
                paragraph_error = "Error: {error} with url: {video_url}".format(error=e, video_url=video_url)
                print(paragraph_error)
                summary_doc.add_paragraph(paragraph_error)
                continue

            doc = docx.Document()
            # add title as heading to summary document
            summary_doc.add_heading(title, 0)
            summary_doc.add_paragraph(paragraph)
            doc.add_heading("Lecture", 0)
            doc.add_paragraph(paragraph)
            doc.save(file_name)

    summary_doc.save("./lectures/summary.docx")
    lectures.close()
    print("Done!")
